import pandas as pd
import numpy as np
import time, locale, os
from datetime import date, datetime
from mailmerge import MailMerge
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.shared import Cm, Pt
from num2words import num2words


pd.set_option("display.max_rows", 500)
pd.set_option("display.max_columns", 500)
pd.set_option("display.width", 150)
locale.setlocale(locale.LC_TIME, "ru_RU")


path = "Y:\\REPORTS\\REPORT2\\" # Path to work folder
template_name = "templ.docx"    # report template file


df = pd.read_excel("Y:\\REPORTS\\REPORT2\\trkn.xlsx", sheet_name="Sheet1") # Excel file with DF for records
dfwho = pd.read_excel("Y:\\REPORTS\\REPORT2\\WorkersData.xls", index_col=0, sheet_name="Sheet4") # Excel file witd DF with workers


def count_empty_blanks(name): 
    """
    here i'm count all empty rows in DF file with records about using blanks
    """
    been = 0
    dftemp = df.loc[df["vid"] == name]
    been = dftemp["tab"].isna().sum()

    return been

#before enter some data
trknbeen = count_empty_blanks("TRKN")
vkldbeen = count_empty_blanks("VKLD")

print(f"TRKN been {trknbeen}, VKLD been {vkldbeen}")


# here i'm count last day in report period and report filename
today = date.today()
last_day = date(today.year, today.month, 1) - relativedelta(days=1)
report_month = last_day.strftime("%B").lower()
report_month_app = last_day.strftime("%m.%Y")
filename = "Report (working with stuff) for " + report_month + " month 2020 [templ].docx"


def who():
    """
    here info about workers and used blanks writing in DF.
    user input tnom, after it pandas find row in df with info about workers and loc this line
    after this from line creating dict. From dict this data inputed in DF and file with info 
    about used blanks.
    In process algorithm checking inputed data and interacting with user.
    """
    try:
        tnom = int(input("Enter №  ")) # here info about workers for find him in DF
        dfwholine = dfwho.loc[dfwho["tnom"] == tnom]
        dfwhodict = dfwholine.to_dict("index")
        info2 = []
        for i in dfwhodict:
                    fio = dfwhodict[i]["famaly"] + " " + dfwhodict[i]["ima"] + " " + dfwhodict[i]["otch"]
                    namecexprof = dfwhodict[i]["namecex"] + "/" + dfwhodict[i]["nameprof"]
                    print(fio, " --- ", namecexprof)
                    info2.extend([dfwhodict[i]["tnom"], fio, namecexprof, report_month_app])

        what = input("Type T or V or CANCEL  ") # here inputed info about what kind of blanks used by worker and write in DF
        if what == "t":
            dftemp = df.loc[df["vid"] == "TRKN"]
            if pd.isna(dftemp["tab"]) is not True:
                for i in range(1):
                    df.iloc[dftemp["tab"].last_valid_index()+1, 2:] = info2
        elif what == "v":
            dftemp = df.loc[df["vid"] == "VKLD"]
            if pd.isna(dftemp["tab"]) is not True:
                for i in range(1):
                    df.iloc[dftemp["tab"].last_valid_index()+1, 2:] = info2
        else:
            print("Check inputed type of blank!")
            who()
    except ValueError:
        print("Check inputed tnom!")
        who()

    df.to_excel("Y:\\REPORTS\\REPORT2\\trkn.xlsx", index=False, sheet_name="Лист1")


end = ""
while end.lower() != "y":
    who()
    end = input("Enter "y"/"Y" to continue ")


def count_what(what):
    """
    here count info about used blanks by type
    """
    when = df.loc[df["kogda"] == report_month_app]
    allblanks = when["tab"].count()
    dftemp = when.loc[when["vid"] == what]
    dtmin = dftemp["num"].min()
    dtmax = dftemp["num"].max()
    dtcount = dftemp["num"].count()
    if what == "TRKN":
        dtpriceallsell = float("%.2f" % (dtcount * 1.29))
        dtostatok = trknbeen - dtcount
    else:
        dtpriceallsell = float("%.2f" % (dtcount * 0.01))
        dtostatok = vkldbeen - dtcount
    print(dtmin, dtmax, dtostatok)

    return allblanks, dtmin, dtmax, dtcount, dtpriceallsell, dtostatok

allblanks, dtmin, dtmax, dtcount, dtpriceallsell, dtostatok = count_what("TRKN")
allblanks, dvmin, dvmax, dvcount, dvpriceallsell, dvostatok = count_what("VKLD")


# here all data about workers and used blanks goes to dict. Info from this dict will be populete to word table 
difortable = df.loc[df["kogda"] == report_month_app].value_counts().to_dict()


def table(name):
    """
    here data from difortable dict populate in table with personal data about workers
    """
    word_document = Document(path + template_name)
    font = word_document.styles["Normal"].font
    font.name = "Times New Roman"
    font.size = Pt(13)
    table0 = word_document.add_table(0, 0)  # we add rows iteratively
    table0.style = word_document.styles["Table Grid"]
    table0.add_column(Cm(3.2))
    table0.add_column(Cm(3.5))
    table0.add_column(Cm(7.3))
    table0.add_column(Cm(2.5))
    table0.add_row()
    hdr_cells = table0.rows[-1].cells
    hdr_cells[0].text = str("№ document")
    hdr_cells[1].text = str("Name")
    hdr_cells[2].text = str("Work unit")
    hdr_cells[3].text = str("Note")
        
    table = word_document.add_table(0, 0)  # we add rows iteratively
    table.style = word_document.styles["Table Grid"]
    table.add_column(Cm(3.2))
    table.add_column(Cm(3.5))
    table.add_column(Cm(7.3))
    table.add_column(Cm(2.5))
    for index, stat_item in enumerate(name.keys()):
        table.add_row()
        vid, num, _, fio, namecexprof, _ = stat_item
        row = table.rows[index]
        row.cells[0].text = str("ПК ") + str(num)
        row.cells[1].text = str(fio)
        row.cells[2].text = str(namecexprof)
        if str(vid) == "TRKN":
            row.cells[3].text = str("employment history book")
        else:
            row.cells[3].text = str("liner in mployment history book")


    word_document.add_paragraph()
    word_document.add_paragraph("Chief of bureau"+" "*88 + "J.A.Smith") #add signature after report
    
    filename = "Report BSR for " + report_month + " month [templ].docx" #temporary report file
    word_document.save(path + filename)
    return filename

table(difortable)

# filling template docx with all another data from my DF
template_1 = path + filename
document = MailMerge(template_1)
document.merge(
    trknbylo = str(trknbeen),
    vkknbylo = str(vkldbeen),
    month2 = str(report_month),
    allblanks = str(allblanks),
    dtmin = str(dtmin),
    dtmax = str(dtmax),
    dtcount = str(dtcount),
    dtcountprice = str(float("%.2f" % (trknbeen*1.29))),
    dtpriceallsell = str(dtpriceallsell),
    dtostatok = str(dtostatok),
    dtostatokprice = str(float("%.2f" % (dtostatok*1.29))),
    dvmin = str(dvmin),
    dvmax = str(dvmax),
    dvcount = str(dvcount),
    dvpriceallsell = str(dvpriceallsell),
    dvostatok = str(dvostatok),
    dvcountprice = str(float("%.2f" % (vkldbeen*0.01))),
    dvostatokprice = str(float("%.2f" % (dvostatok*0.01))),
    propis = str(num2words(allblanks))
)
filename = "Report BSR for " + report_month + " " + date.strftime(date.today(), "%Y") + ".docx"
document.write(path + filename)  # save file to folder
document.close() # close document

# remove temporary file with report
os.remove(path + ("Report BSR for " + report_month + " month [templ].docx"))


print(filename)
input("Press ENTER to exit from app...")
os.startfile(path)